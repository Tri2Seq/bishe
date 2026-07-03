"""严格按照中期进展报告模板生成文档"""
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ===== 页面设置 A4 =====
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.17)
section.right_margin = Cm(3.17)

# 默认段落样式
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
pf = style.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
pf.space_after = Pt(0)
pf.space_before = Pt(0)

def add_centered(text, size=12, bold=False, space_before=0, space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_info_line(label, value, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f'{label}：{value}')
    run.font.size = Pt(size)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_body(text, indent=True, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)  # 两字符缩进
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_section_title(text, size=14):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = True
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return p

# ================================================================
# 封面页
# ================================================================
add_centered('', size=12)
add_centered('学    院', size=14)
add_centered('', size=12)
add_centered('编    号', size=14)
add_centered('', size=12)
add_centered('', size=12)
add_centered('', size=12)
add_centered('硕士研究生', size=26, bold=True, space_before=24)
add_centered('中期进展报告', size=26, bold=True, space_after=24)
add_centered('', size=12)
add_centered('', size=12)

add_info_line('学    号', '3124354083')
add_info_line('研 究 生', '施佳希')
add_info_line('导    师', '胡建晨 教授')
add_info_line('论文题目', '面向多楼层场景的多机器人协同配送')
add_info_line('        ', '与计算资源调度研究')
add_info_line('学    科', '控制工程')
add_info_line('填写时间', '2026 年 7 月')

doc.add_page_break()

# ================================================================
# 填写说明页
# ================================================================
add_section_title('填  写  说  明', size=16)
doc.add_paragraph('')
add_body('一、硕士研究生中期进展报告必须采用A4纸双面打印，左侧装订成册，各栏空格不够时，请自行加页。', indent=False)
add_body('二、中期进展报告正文部分字数为4000-5000字。', indent=False)
add_body('三、硕士研究生中期进展报告通过后，本报告由各系存档。', indent=False)

doc.add_page_break()

# ================================================================
# 正文
# ================================================================

# ----- 研究题目 -----
add_section_title('研究题目')
add_body('面向多楼层场景的多机器人协同配送与计算资源调度研究', indent=False)

# ----- 一、研究内容简介 -----
add_section_title('一、研究内容简介（300～500字）')

add_body(
    '随着智能物流与楼宇自动化技术的发展，多机器人跨楼层配送逐渐成为高层建筑物流的核心需求。'
    '电商仓库、医院、写字楼等多楼层场景中，配送机器人需要借助电梯在不同楼层间运送货物，'
    '同时还需要云端GPU进行单目深度估计推理以辅助电梯按钮操控。'
    '该场景涉及三类稀缺资源的竞争调度——电梯通道、异构机器人车队和云端GPU计算资源——构成了一个多约束、多目标的复杂优化问题。'
)

add_body(
    '本研究将上述问题建模为两个串联模块：模块一为多机器人多楼层配送调度，'
    '核心决策包括任务-机器人分配、执行序列优化、电梯群控交互及多机器人中继协作，'
    '以最小化完工时间、总能耗和加权延迟为三目标；'
    '模块二为GPU资源调度，以模块一产出的机器人动作时间线为输入，'
    '优化GPU分配、模型预加载与显存碎片管理。'
    '在算法层面，本研究提出了NSGA-II-EBO（Elevator-Bottleneck Optimization）混合优化算法，'
    '利用问题的电梯瓶颈数学结构实现有效搜索空间缩减；'
    '在仿真层面，构建了基于电梯群控系统的离散时间仿真器，支持24小时周期性任务流的端到端仿真评估。'
)

# ----- 二、研究工作进展 -----
add_section_title('二、研究工作进展')

# (1) 系统建模
add_body(
    '（1）系统建模与问题定义。'
    '完成了完整的多楼层配送场景数学建模。物理环境建模为F层建筑（5-20层），每层设有取货站和送货站，'
    '楼层间仅通过电梯连通。底层设为收货区，中层为存储区，高层为打包发货区，'
    '各功能楼层间的物流需求构成了跨楼层配送的主要驱动力。楼层内站点间距设定为30米，站点到电梯大厅距离为50米。'
    '机器人建模为三种异构类型——重载型（载重80kg、速度0.8m/s、续航3000m、能耗5.0J/m）、'
    '标准型（载重30kg、速度1.5m/s、续航2500m、能耗3.0J/m）、快递型（载重10kg、速度2.5m/s、续航2000m、能耗2.0J/m），'
    '车队组成比例为30%:50%:20%。电梯分为货梯（容量2台机器人、速度0.4m/s、开门时间8s）'
    '和客梯（容量1台、速度0.6m/s、开门时间5s），货客比例为4:6。'
    '配送任务建模为六元组（起点楼层、终点楼层、重量、优先级、截止时间、最早开始时间），'
    '优先级分为紧急（权重10）、普通（权重3）、批量（权重1）三级，跨楼层任务占比约64%。'
    '优化目标定义为三个指标的同时最小化：最大完工时间（makespan）、总能耗（机器人+电梯）和加权延迟（优先级权重×超期时间）。'
    '约束条件涵盖任务唯一性、机器人载重上限、续航限制、电梯容量及时间窗等。'
)

# (2) 电梯群控
add_body(
    '（2）电梯群控系统建模。'
    '区别于现有工作中机器人直接选择电梯的简化假设，本研究建模了真实的电梯群控系统（EGCS）。'
    '在实际建筑中，机器人按下楼层呼叫按钮后，由楼宇中央群控系统根据派车算法决定派哪一部电梯响应。'
    '本研究实现了三种经典派车算法：ETA算法基于预估到达时间最短原则，综合考虑直达距离、中间停站数、开关门时间及满载惩罚；'
    'SCAN算法令电梯沿单一方向扫描到头再折返；NC算法选择距呼叫楼层距离最近的空闲电梯。'
    '每部电梯建模为独立状态机（IDLE→MOVING_TO_CALL→DOOR_OPENING→MOVING_TO_DEST），'
    '支持多乘客共乘、满载判定和方向优先。能耗区分空载（0.008 kWh/层）和满载（0.015 kWh/层）。'
    '实验表明ETA算法在任务完成率上最优（85%），比NC和SCAN高出约10个百分点。'
)

# (3) MILP
add_body(
    '（3）MILP精确求解模型。'
    '建立了包含12类约束的混合整数线性规划模型，核心创新在于电梯不相交约束（Disjunctive Scheduling）：'
    '将每部电梯视为不可抢占的机器，引入二元变量σ[i,i\',e]表示同一电梯上两次使用的先后顺序，'
    '约束中包含重定位时间（电梯空驶到下一次服务起点的时间）。'
    '该建模方式精确刻画了电梯资源的时序竞争，为启发式算法提供最优性下界参考。'
    'MILP模型还包含任务唯一性、载重、续航、时间窗和电梯容量等约束，适用于20个任务以内的小规模实例精确求解。'
)

# (4) NSGA-II-EBO
add_body(
    '（4）NSGA-II-EBO混合优化算法。'
    '这是本研究的核心算法贡献。通过问题特征分析发现三个关键结构性质：'
    '电梯是系统瓶颈（电梯服务时间下界240s > 机器人下界159s）；'
    '同楼层任务存在电梯通道竞争和重载任务存在稀缺机器人竞争；'
    '50kg以上关键任务仅4/15机器人可执行，其分配决定全局质量。'
    '基于上述分析，算法采用双层架构：上层为NSGA-II进化搜索，'
    '决策变量包括任务-机器人分配、中继楼层、中继机器人及错峰延迟，搜索空间O(M^N)；'
    '下层为结构感知的序列优化器，采用权重感知启发式规则（EDD、最近邻、紧急优先）与2-opt局部搜索，'
    '时间复杂度O(N log N)。相比标准NSGA-II的O(M^N × N!)搜索空间，NSGA-II-EBO通过分解大幅缩减了有效搜索空间。'
    '下层优化器根据三目标权重自适应选择策略：高Makespan权重下优先处理紧急任务，'
    '高能耗权重下采用最近邻启发式最小化空跑距离，高延迟权重下按截止时间排序。'
    '每次评估尝试三个权重方向，取快速估算最优者送入仿真器精确评估。'
)

add_body(
    '算法设计了四类专用遗传算子：冲突感知交叉（关键任务从更优父代继承，普通任务均匀交叉）、'
    '负载均衡变异（从最忙机器人转移任务到最闲机器人）、'
    '中继变异（启用/关闭长途任务中继、调整中继楼层或更换协作机器人）、'
    '策略参数变异（错峰延迟参数协同进化）。'
    '种群初始化采用三策略混合：30%关键任务优先（利用冲突图避免高冲突对同分配）、'
    '20%楼层聚类（空间局部性）、50%随机（多样性保证）。'
    '问题分析模块为算法设计提供定量依据：电梯瓶颈分析给出Makespan下界，'
    '任务冲突图构建N×N矩阵量化楼层、电梯、机器人和时间窗四维竞争强度，'
    '关键任务识别综合稀缺性、紧迫性、占用度和优先级四因子排序。'
)

# (5) 中继
add_body(
    '（5）多机器人中继协作配送。'
    '针对长途跨楼层任务（楼层差≥3），设计了多机器人中继协作机制。'
    '一个长途任务可在中继楼层由第一台机器人交接给第二台，第一台释放后立即处理其他任务，'
    '减少单台机器人长时间占用、提高车队整体利用率。'
    '中继机制在数据层新增TaskSegment和RelayPlan两个结构；'
    '仿真层新增RELAY_DROPOFF阶段和relay_buffer缓冲区，支持前驱约束检查；'
    '优化层将中继决策编码为进化基因（gene_relay和gene_relay_robot），'
    '由NSGA-II-EBO自动学习何时中继、中继楼层选择和协作机器人匹配。'
    '修复机制保证中继合法性：中继楼层严格位于起终点之间，第二段机器人不同于第一段且满足载重要求。'
    '实验表明大规模场景下中继优化使Makespan进一步降低31.4%。'
)

# (6) 批量+24h
add_body(
    '（6）批量携带与24小时周期建模。'
    '实现了机器人批量携带——同楼层多个待取任务在载重限制内一次取走，乘一次电梯沿途逐层送达，'
    '显著减少电梯使用次数。基于物流园区运营模式建立24小时周期性任务生成模型，'
    '定义6个时段：早班启动（06-08时，8任务/h）、上午高峰（08-12时，30任务/h，上行60%）、'
    '午间低谷（12-13时，5任务/h）、下午高峰（13-17时，28任务/h，下行60%）、'
    '晚间收尾（17-20时，12任务/h）、夜间低谷（20-06时，3任务/h）。'
    '日均生成约319个任务，峰谷比8.2倍，真实反映了物流园区的周期性运营特征。'
)

# (7) GPU
add_body(
    '（7）GPU资源调度（模块二）。'
    '模块二以模块一的机器人动作时间线为输入，提取GPU深度估计推理需求段'
    '（接近电梯8s、按按钮7s、出电梯5s）和空闲窗口（乘电梯期间不需GPU）。'
    '提出预测性GPU调度策略（PredictiveScheduler），利用24小时任务周期的已知模式进行三层决策：'
    '预加载（需求前10s提前加载模型避免冷启动）、智能保持（间隔短时保持模型常驻避免反复加载）、'
    '预释放（低谷期主动释放减少碎片）。'
    '对比实现了FragAware（碎片感知）、OnDemand（按需分配）、Fixed（固定绑定）和RoundRobin（轮询）等基线。'
)

# (8) 仿真器
add_body(
    '（8）仿真器开发。'
    '构建了基于电梯群控系统的离散时间仿真器SimulatorV2（步长0.5s），核心特性：'
    '电梯由群控系统统一调度（支持ETA/SCAN/NC）；'
    '机器人12种运行状态，RIDING和WAITING阶段事件驱动以避免死锁；'
    '支持批量取货、多站送货和中继交接；自动标注GPU需求时段；支持24小时连续仿真。'
)

# (9) 实验
add_body(
    '（9）实验结果。'
    '算法对比实验：三个规模场景（小型15任务、中型30任务、大型50任务）各5次独立运行。'
    '大型场景（15层、12机器人、5电梯、50任务）结果：'
    'Greedy平均Makespan 1003±73秒，EBO 342±69秒（降低65.9%），'
    'EBO+Relay 235±54秒（降低76.6%，中继额外降低31.4%）。'
    '能耗方面EBO+Relay均值10633J，相比Greedy的44207J降低76.0%。'
    '所有EBO方案加权延迟均降至零。EBO+Relay平均启用5.8个中继任务。'
    '中型场景EBO+Relay Makespan 214±36秒（Greedy 685±72秒，降低68.7%）。'
)

add_body(
    '派车算法对比：ETA的Makespan 318秒（平均等待3.1秒）最优，'
    'SCAN 354秒（4.3秒），NC 470秒（5.5秒）。'
    '24小时全链路实验（15层、20机器人、4电梯、319任务）：'
    'ETA+Greedy完成257/319（80.6%），ETA+EBO+Relay完成279/319（87.5%）；'
    'NC+EBO+Relay 299/319（93.7%，提升26.7%），SCAN+EBO+Relay 300/319（94.0%，提升27.1%）。'
    'GPU调度对比（紧张配置2GPU/4并发）：'
    'Predictive策略等待0秒、冷启动0次、碎片率0.026；'
    'FragAware等待0秒、冷启动683次；OnDemand等待2394秒、碎片率0.937。'
    'Predictive通过预加载和智能保持实现零等待零冷启动。'
)

# ----- 三、下一步的工作计划 -----
add_section_title('三、下一步的工作计划')

add_body(
    '基于当前已完成的系统建模、算法设计和仿真实验，下一步工作计划如下：'
)

add_body(
    '（1）单目深度估计模型研发。'
    '开发基于VAE+DiT（Diffusion Transformer）架构的轻量化单目深度估计模型，'
    '实现对电梯按钮区域的精确深度感知。模型需满足边缘部署推理速度要求（<100ms），'
    '并在云-边协作框架下支持模型的动态加载与卸载。'
)

add_body(
    '（2）云-边协同推理框架。'
    '设计机器人端与云端GPU之间的推理任务协同机制，包括推理请求调度、模型版本管理、网络延迟补偿等。'
    '将单目深度估计的推理负载集成到GPU预测性调度框架中，实现端到端闭环验证。'
)

add_body(
    '（3）算法理论分析。'
    '对NSGA-II-EBO算法进行收敛性和计算复杂度的理论分析，'
    '推导电梯瓶颈下界的紧致性，分析中继机制对搜索空间和求解质量的影响。'
)

add_body(
    '（4）大规模实验与对比。'
    '在更大规模场景（30层、50机器人、100+任务）下进行系统性实验，'
    '与更多对比方法（遗传算法GA、粒子群PSO、强化学习RL）进行统计检验。'
    '增加Wilcoxon秩和检验等假设检验以验证显著性。'
)

add_body(
    '（5）论文撰写与投稿。'
    '整理研究成果撰写学位论文初稿。'
    '计划将配送调度模块的核心算法投稿至控制与优化领域会议或期刊。'
)

# ----- 四、导师评语 -----
add_section_title('四、导师评语')

for _ in range(12):
    doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('\t签名：\t\t\t日期：      年   月   日')
run.font.size = Pt(12)

# ===== 保存 =====
output_path = '/home/ubuntu/Documents/sjx/gra/midterm.docx'
doc.save(output_path)
print(f'Saved to {output_path}')

# 字数统计
import re
full_text = '\n'.join([p.text for p in doc.paragraphs])
chinese_chars = len(re.findall(r'[一-鿿]', full_text))
all_chars = len(re.sub(r'\s', '', full_text))
print(f'中文字数: {chinese_chars}')
print(f'总字符数(含英文数字): {all_chars}')
